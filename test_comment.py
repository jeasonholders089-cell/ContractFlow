"""
测试批注生成功能
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from backend.utils.comment_generator import CommentGenerator

# 创建一个简单的测试文档
test_doc_path = "test_contract.docx"
doc = Document()
doc.add_paragraph("这是一个测试合同文档。")
doc.add_paragraph("甲方：测试公司")
doc.add_paragraph("乙方：另一方公司")
doc.add_paragraph("本合同条款不明确，需要修改。")
doc.add_paragraph("违约责任未明确规定。")
doc.save(test_doc_path)

print(f"创建测试文档: {test_doc_path}")

# 测试批注添加
comment_gen = CommentGenerator(test_doc_path)

# 在第3段添加批注（索引从0开始，所以是2）
success = comment_gen.add_comment(2, "【高风险】条款不明确\n建议：应明确具体条款内容", "AI审核")
print(f"添加第1个批注: {'成功' if success else '失败'}")

# 在第4段添加批注
success = comment_gen.add_comment(3, "【中风险】违约责任未明确\n建议：增加违约金计算方式", "AI审核")
print(f"添加第2个批注: {'成功' if success else '失败'}")

# 保存带批注的文档
output_path = "test_contract_with_comments.docx"
comment_gen.save(output_path)
print(f"保存带批注的文档: {output_path}")
print("\n请在 Word 或 WPS 中打开该文档，查看批注是否显示")
print("如果批注没有显示，会看到内联的红色文字标记")
