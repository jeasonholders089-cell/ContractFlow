"""
批注生成模块
"""
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, List
from xml.etree import ElementTree as ET
import re
import os
import tempfile
import zipfile


class CommentGenerator:
    """批注生成器 - 完整实现Word批注功能"""

    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self._comment_id_start = 0
        self._ensure_comments_part()

    def _ensure_comments_part(self):
        """确保文档有批注部分"""
        # 批注功能会在实际添加批注时自动创建
        # 这里不需要预先检查
        pass

    def add_comment(self, paragraph_index: int, text: str, author: str = "AI审核") -> bool:
        """
        为指定段落添加批注

        Args:
            paragraph_index: 段落索引
            text: 批注内容
            author: 批注作者

        Returns:
            是否成功添加
        """
        if paragraph_index >= len(self.doc.paragraphs):
            return False

        paragraph = self.doc.paragraphs[paragraph_index]

        # 如果段落为空，跳过
        if not paragraph.text.strip():
            return False

        # 先尝试XML批注，失败则使用备用方案
        try:
            if len(paragraph.runs) == 0:
                run = paragraph.add_run(paragraph.text)
                paragraph.clear()
                paragraph.add_run(run.text)
                paragraph = self.doc.paragraphs[paragraph_index]

            # 尝试添加XML批注
            self._add_comment_to_paragraph(paragraph, text, author)
        except Exception as e:
            print(f"XML批注失败，使用备用方案: {e}")
            # 使用备用方案
            self._add_comment_inline(paragraph, text, author)

        return True

    def _add_comment_to_paragraph(self, paragraph, text: str, author: str):
        """为段落添加批注 - 使用更可靠的XML方式"""
        p_element = paragraph._element

        # 找到所有 w:r (run) 元素
        run_elements = p_element.findall(qn("w:r"), p_element.nsmap)
        if not run_elements or len(run_elements) == 0:
            raise Exception("没有找到run元素")

        # 选择第一个run来添加批注
        target_run = run_elements[0]

        # 为目标run添加背景色高亮
        # 获取或创建run的rPr属性
        run_props = target_run.find(qn("w:rPr"), target_run.nsmap)
        if run_props is None:
            run_props = OxmlElement("w:rPr")
            target_run.insert(0, run_props)
        else:
            # 清空现有属性
            for child in list(run_props):
                run_props.remove(child)

        # 添加背景色为黄色
        # 使用shd元素添加底纹
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFFF00")  # 黄色背景
        run_props.append(shd)

        # 创建批注范围开始标记
        comment_start = OxmlElement("w:commentRangeStart")
        comment_start.set(qn("w:id"), str(self._comment_id_start))

        # 创建批注范围结束标记
        comment_end = OxmlElement("w:commentRangeEnd")
        comment_end.set(qn("w:id"), str(self._comment_id_start))

        # 创建批注引用run
        comment_ref_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "CommentReference")
        r_pr.append(r_style)
        comment_ref_run.append(r_pr)

        comment_ref = OxmlElement("w:annotationRef")
        comment_ref.set(qn("w:id"), str(self._comment_id_start))
        comment_ref_run.append(comment_ref)

        # 找到目标run的索引位置
        run_index = list(p_element).index(target_run)

        # 正确的插入顺序：
        # 1. 在目标run之前插入commentRangeStart
        p_element.insert(run_index, comment_start)

        # 2. 插入批注引用run（在目标run之后）
        # 由于插入了comment_start，索引需要+1
        p_element.insert(run_index + 2, comment_ref_run)

        # 3. 插入commentRangeEnd（在批注引用之后）
        p_element.insert(run_index + 3, comment_end)

        # 添加批注内容到文档 - 如果这失败，抛出异常
        self._add_comment_to_document(self._comment_id_start, author, text)

        self._comment_id_start += 1

        print(f"✓ 批注标记已添加 (ID: {self._comment_id_start - 1})")

    def _add_comment_to_document(self, comment_id: int, author: str, text: str):
        """将批注添加到文档的批注部分 - 使用正确的方法"""
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn

            # 获取文档的main part
            doc_part = self.doc.part

            # 尝试获取现有的批注部分
            comments_part = None

            # 遍历关系查找批注部分
            for rel in doc_part.rels.values():
                if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments":
                    comments_part = rel.target_part
                    break

            # 如果批注部分不存在，创建新的
            if comments_part is None:
                # 创建新的批注部分 - 使用正确的方法
                # 添加到document.xml.rels中
                rel_id = doc_part.rels.add(
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
                )

                # 创建comments.xml内容
                comments_xml = '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
                comments_element = parse_xml(comments_xml)

                # 创建新的Part
                from docx.opc.packuri import PackURI
                comments_uri = PackURI("/word/comments.xml")

                # 使用doc_part.package来添加part
                comments_part = doc_part.package.part_related_by[rel_id]

                # 设置element
                if hasattr(comments_part, '_element'):
                    comments_part._element = comments_element
                else:
                    # 创建part并设置内容
                    doc_part.package.save_part(comments_uri, comments_element.xml)

            # 获取comments元素
            comments_element = comments_part.element

            # 创建批注
            comment = OxmlElement("w:comment")
            comment.set(qn("w:id"), str(comment_id))
            comment.set(qn("w:author"), author)
            # 添加初始日期
            from datetime import datetime
            comment.set(qn("w:date"), datetime.now().strftime("%Y-%m-%dT%H:%M:00Z"))

            # 创建批注段落
            p = OxmlElement("w:p")

            # 创建批注内容的run
            r = OxmlElement("w:r")

            # 添加文本
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)

            p.append(r)
            comment.append(p)

            # 将批注添加到comments元素
            comments_element.append(comment)

            print(f"批注内容已添加到comments部分: {text[:50]}...")

        except Exception as e:
            # 如果无法添加到批注部分，重新抛出异常
            import traceback
            print(f"批注内容添加失败: {e}")
            traceback.print_exc()
            raise e

    def _add_comment_inline(self, paragraph, text: str, author: str):
        """备用方案：在段落后添加批注标记"""
        try:
            print(f"[备用方案] 开始添加批注...")
            # 直接在段落的XML元素中添加run
            p_element = paragraph._element

            # 创建新的run元素
            new_run = OxmlElement("w:r")

            # 创建run属性
            r_pr = OxmlElement("w:rPr")

            # 设置颜色为红色
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "FF0000")
            r_pr.append(color)

            # 设置字体大小
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")  # 10pt = 20 half-points
            r_pr.append(sz)

            new_run.append(r_pr)

            # 创建文本元素（不使用emoji）
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            comment_text = f"\n[{author}: {text}]"
            t.text = comment_text
            new_run.append(t)

            # 添加到段落末尾
            p_element.append(new_run)
            print(f"[备用方案] 完成")

        except Exception as e:
            print(f"[备用方案] 失败: {e}")
            import traceback
            traceback.print_exc()

    def add_review_summary(self, issues: List[Dict]) -> bool:
        """
        在文档开头添加审核摘要

        Args:
            issues: 问题列表

        Returns:
            是否成功添加
        """
        try:
            # 在文档开头插入审核摘要
            summary_para = self.doc.paragraphs[0]._element
            new_para = summary_para.addprevious(
                self._create_summary_element(issues)
            )
            return True
        except Exception:
            return False

    def _create_summary_element(self, issues: List[Dict]):
        """创建摘要元素"""
        # 创建新的段落元素
        p = OxmlElement("w:p")
        p.set(qn("w:rPr"), "")

        # 添加标题
        title_r = OxmlElement("w:r")
        title_t = OxmlElement("w:t")
        title_t.text = "=== AI审查报告 ==="
        title_r.append(title_t)
        p.append(title_r)

        # 添加问题汇总
        for i, issue in enumerate(issues[:10], 1):  # 最多显示10条
            # 创建运行
            r = OxmlElement("w:r")

            # 创建文本
            t = OxmlElement("w:t")
            t.text = f"\n{i}. [{issue.get('severity', '中')}] {issue.get('category', '')}: {issue.get('problem', '')[:50]}"
            r.append(t)
            p.append(r)

        return p

    def save(self, output_path: str):
        """保存文档"""
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        self.doc.save(output_path)

    def add_issues_as_comments(self, issues_with_location: List[Dict]) -> int:
        """
        批量添加问题批注 - 使用内联红色文字方式
        """
        print("使用内联红色文字批注方式...")
        count = 0
        for item in issues_with_location:
            issue = item.get("issue", {})
            location = item.get("location", {})

            if location:
                paragraph_index = location.get("index", -1)
                if paragraph_index >= 0:
                    comment_text = (
                        f"【{issue.get('severity', '中')}风险】{issue.get('problem', '')}\n"
                        f"建议：{issue.get('suggestion', '')}"
                    )
                    if self.add_comment(paragraph_index, comment_text):
                        count += 1
        return count

    def create_review_report(self, issues: List[Dict], summary: str) -> str:
        """
        创建审查报告（纯文本）

        Args:
            issues: 问题列表
            summary: 审查摘要

        Returns:
            报告文本
        """
        report_lines = [
            "=" * 60,
            "合同审查报告",
            "=" * 60,
            "",
            summary,
            "",
            "=" * 60,
            "问题详情",
            "=" * 60,
            ""
        ]

        for i, issue in enumerate(issues, 1):
            report_lines.extend([
                f"{i}. {issue.get('category', '')} - {issue.get('severity', '')}风险",
                f"   位置：{issue.get('location_hint', '')}",
                f"   原文：{issue.get('original_text', '')}",
                f"   问题：{issue.get('problem', '')}",
                f"   建议：{issue.get('suggestion', '')}",
                ""
            ])

        return "\n".join(report_lines)
