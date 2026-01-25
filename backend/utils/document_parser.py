"""
Word 文档解析模块
"""
from docx import Document
from typing import List, Dict, Any
import re


class DocumentParser:
    """文档解析器"""

    def __init__(self):
        pass

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        解析 Word 文档

        Returns:
            {
                "paragraphs": [(text, style_name, index), ...],
                "tables": [[row1, row2, ...], ...],
                "full_text": str,
                "structure": {...}
            }
        """
        doc = Document(file_path)

        # 解析段落
        paragraphs = []
        for idx, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else "Normal"
            paragraphs.append({
                "text": para.text,
                "style": style_name,
                "index": idx
            })

        # 解析表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # 生成全文
        full_text = "\n".join([p["text"] for p in paragraphs if p["text"].strip()])

        # 检测文档结构
        structure = self._detect_structure(paragraphs)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": full_text,
            "structure": structure,
            "total_paragraphs": len(paragraphs)
        }

    def _detect_structure(self, paragraphs: List[Dict]) -> Dict[str, Any]:
        """检测文档结构"""
        sections = []
        current_section = None
        section_num = 0

        for para in paragraphs:
            text = para["text"].strip()
            style = para["style"]

            # 检测是否为标题
            if self._is_heading(style) or self._is_numbered_title(text):
                section_num += 1
                current_section = {
                    "number": section_num,
                    "title": text,
                    "start_index": para["index"]
                }
                sections.append(current_section)
            elif current_section:
                current_section["end_index"] = para["index"]

        return {
            "sections": sections,
            "total_sections": section_num
        }

    def _is_heading(self, style_name: str) -> bool:
        """判断是否为标题样式"""
        if style_name is None:
            return False
        heading_patterns = ["Heading", "标题", "Title", "标题 1", "Heading 1"]
        return any(pattern in style_name for pattern in heading_patterns)

    def _is_numbered_title(self, text: str) -> bool:
        """判断是否为编号标题"""
        # 匹配 "第X条"、"1."、"一、" 等模式
        patterns = [
            r"^第[一二三四五六七八九十百]+条",
            r"^\d+\.",
            r"^[一二三四五六七八九十]+、",
        ]
        return any(re.match(pattern, text) for pattern in patterns)

    def estimate_tokens(self, text: str) -> int:
        """估算 Token 数量"""
        # 中文约 1.5 字符 = 1 token，英文约 4 字符 = 1 token
        chinese_chars = len([c for c in text if '\u4e00' <= c <= '\u9fff'])
        other_chars = len(text) - chinese_chars
        return int(chinese_chars / 1.5 + other_chars / 4)

    def should_split(self, parsed_doc: Dict, max_tokens: int) -> bool:
        """判断是否需要分段处理"""
        estimated_tokens = self.estimate_tokens(parsed_doc["full_text"])
        return estimated_tokens > max_tokens

    def split_by_sections(self, parsed_doc: Dict, max_tokens: int) -> List[Dict]:
        """按章节分段"""
        sections = parsed_doc["structure"]["sections"]
        paragraphs = parsed_doc["paragraphs"]

        if not sections:
            # 没有章节结构，按段落数量分段
            return self._split_by_paragraphs(paragraphs, max_tokens)

        # 按章节分段
        result = []
        current_section = None
        current_paras = []
        current_tokens = 0

        for para in paragraphs:
            # 检查是否是章节开始
            section = next((s for s in sections if s["start_index"] == para["index"]), None)

            if section:
                # 保存当前分段
                if current_paras:
                    result.append({
                        "section_number": len(result) + 1,
                        "paragraphs": current_paras,
                        "text": "\n".join([p["text"] for p in current_paras])
                    })

                current_section = section
                current_paras = [para]
                current_tokens = self.estimate_tokens(para["text"])
            else:
                para_tokens = self.estimate_tokens(para["text"])

                # 检查是否超过限制
                if current_tokens + para_tokens > max_tokens and current_paras:
                    result.append({
                        "section_number": len(result) + 1,
                        "paragraphs": current_paras,
                        "text": "\n".join([p["text"] for p in current_paras])
                    })
                    current_paras = [para]
                    current_tokens = para_tokens
                else:
                    current_paras.append(para)
                    current_tokens += para_tokens

        # 保存最后一段
        if current_paras:
            result.append({
                "section_number": len(result) + 1,
                "paragraphs": current_paras,
                "text": "\n".join([p["text"] for p in current_paras])
            })

        return result

    def _split_by_paragraphs(self, paragraphs: List[Dict], max_tokens: int) -> List[Dict]:
        """按段落分段（无章节结构时）"""
        result = []
        current_paras = []
        current_tokens = 0

        for para in paragraphs:
            para_tokens = self.estimate_tokens(para["text"])

            if current_tokens + para_tokens > max_tokens and current_paras:
                result.append({
                    "section_number": len(result) + 1,
                    "paragraphs": current_paras,
                    "text": "\n".join([p["text"] for p in current_paras])
                })
                current_paras = [para]
                current_tokens = para_tokens
            else:
                current_paras.append(para)
                current_tokens += para_tokens

        if current_paras:
            result.append({
                "section_number": len(result) + 1,
                "paragraphs": current_paras,
                "text": "\n".join([p["text"] for p in current_paras])
            })

        return result
