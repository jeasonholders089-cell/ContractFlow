"""
合同文档生成器
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, Any
import os
import re
from datetime import datetime


class ContractDocumentBuilder:
    """合同文档生成器 - 生成格式化的 Word 合同文档"""

    def __init__(self):
        """初始化文档生成器"""
        self.doc = None

    def create_contract_document(
        self,
        content: str,
        title: str,
        metadata: Dict[str, Any] = None,
        output_path: str = None
    ) -> str:
        """
        创建格式化的合同文档

        Args:
            content: 合同内容（纯文本）
            title: 合同标题
            metadata: 元数据（可选）
            output_path: 输出路径（可选，如果不提供则自动生成）

        Returns:
            生成的文档路径
        """
        # 创建新文档
        self.doc = Document()

        # 设置文档样式
        self._setup_document_styles()

        # 添加文档元数据
        if metadata:
            self._add_metadata(metadata)

        # 添加合同标题
        self._add_title(title)

        # 添加合同内容
        self._add_content(content)

        # 添加签署栏
        self._add_signature_section()

        # 保存文档
        if not output_path:
            output_path = self._generate_output_path(title)

        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        # 保存文档
        self.doc.save(output_path)

        return output_path

    def _setup_document_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_metadata(self, metadata: Dict[str, Any]):
        """添加文档元数据"""
        core_properties = self.doc.core_properties
        core_properties.author = metadata.get('author', 'ContractFlow AI')
        core_properties.title = metadata.get('title', '')
        core_properties.subject = metadata.get('subject', '合同文档')
        core_properties.created = datetime.utcnow()

    def _add_title(self, title: str):
        """添加合同标题"""
        # 添加标题段落
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加标题文本
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run.font.color.rgb = RGBColor(0, 0, 0)

        # 设置中���字体
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加空行
        self.doc.add_paragraph()

    def _add_content(self, content: str):
        """
        添加合同内容

        解析合同文本并格式化：
        - 识别条款编号（第一条、第二条等）
        - 识别甲乙方信息
        - 保持段落结构
        """
        # 按行分割内容
        lines = content.split('\n')

        for line in lines:
            line = line.strip()

            # 跳过空行
            if not line:
                self.doc.add_paragraph()
                continue

            # 识别条款标题（第一条、第二条等）
            if self._is_clause_title(line):
                self._add_clause_title(line)
            # 识别甲乙方信息
            elif self._is_party_info(line):
                self._add_party_info(line)
            # 普通段落
            else:
                self._add_normal_paragraph(line)

    def _is_clause_title(self, line: str) -> bool:
        """判断是否为条款标题"""
        # 匹配：第一条、第二条、第1条、第2条等
        patterns = [
            r'^第[一二三四五六七八九十百]+条',
            r'^第\d+条',
            r'^\d+\.',
            r'^[一二三四五六七八九十]+、'
        ]
        for pattern in patterns:
            if re.match(pattern, line):
                return True
        return False

    def _is_party_info(self, line: str) -> bool:
        """判断是否为甲乙方信息"""
        return line.startswith('甲方') or line.startswith('乙方') or \
               line.startswith('出租方') or line.startswith('承租方') or \
               line.startswith('委托方') or line.startswith('受托方')

    def _add_clause_title(self, line: str):
        """添加条款标题"""
        para = self.doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 设置段落间距
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)

    def _add_party_info(self, line: str):
        """添加甲乙方信息"""
        para = self.doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置段落间距
        para.paragraph_format.space_after = Pt(3)

    def _add_normal_paragraph(self, line: str):
        """添加普通段落"""
        para = self.doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置段落格式
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进2字符

    def _add_signature_section(self):
        """添加签署栏"""
        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 添加签署栏标题
        para = self.doc.add_paragraph()
        run = para.add_run('（以下无正文）')
        run.font.size = Pt(12)
        run.font.name = '宋体'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 添加甲方签署栏
        self._add_signature_block('甲方（盖章）：', '日期：    年    月    日')

        # 添加空行
        self.doc.add_paragraph()

        # 添加乙方签署栏
        self._add_signature_block('乙方（盖章）：', '日期：    年    月    日')

    def _add_signature_block(self, party_label: str, date_label: str):
        """添加签署块"""
        # 甲方/乙方标签
        para1 = self.doc.add_paragraph()
        run1 = para1.add_run(party_label)
        run1.font.size = Pt(12)
        run1.font.name = '宋体'

        # 添加空行用于签名
        self.doc.add_paragraph()

        # 日期标签
        para2 = self.doc.add_paragraph()
        run2 = para2.add_run(date_label)
        run2.font.size = Pt(12)
        run2.font.name = '宋体'

    def _generate_output_path(self, title: str) -> str:
        """生成输出路径"""
        # 清理标题中的非法字符
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)

        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{safe_title}_{timestamp}.docx"

        # 使用 storage/drafts 目录
        output_dir = os.path.join('storage', 'drafts')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        return os.path.join(output_dir, filename)


def create_contract_document(
    content: str,
    title: str,
    metadata: Dict[str, Any] = None,
    output_path: str = None
) -> str:
    """
    便捷函数：创建合同文档

    Args:
        content: 合同内容
        title: 合同标题
        metadata: 元数据
        output_path: 输出路径

    Returns:
        生成的文档路径
    """
    builder = ContractDocumentBuilder()
    return builder.create_contract_document(content, title, metadata, output_path)

