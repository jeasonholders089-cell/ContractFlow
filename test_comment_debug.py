"""
详细调试批注生成
"""
import sys
sys.path.insert(0, r'f:\Agent项目\AI合同审查')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 创建测试文档
doc = Document()
para = doc.add_paragraph('测试段落')
doc.save('test_debug2.docx')
print(f'原始文档: {len(doc.paragraphs)} 段落')

# 重新加载并添加批注
doc2 = Document('test_debug2.docx')
paragraph = doc2.paragraphs[0]
print(f'段落文本: "{paragraph.text}"')
print(f'初始runs数: {len(paragraph.runs)}')

# 手动添加一个带颜色的run
p_element = paragraph._element
print(f'段落XML元素: {p_element.tag}')

# 创建新的run元素
new_run = OxmlElement("w:r")

# 创建run属性
r_pr = OxmlElement("w:rPr")

# 设置颜色
color = OxmlElement("w:color")
color.set(qn("w:val"), "FF0000")
r_pr.append(color)

# 设置字体大小
sz = OxmlElement("w:sz")
sz.set(qn("w:val"), "20")
r_pr.append(sz)

new_run.append(r_pr)

# 创建文本元素
t = OxmlElement("w:t")
t.set(qn("xml:space"), "preserve")
t.text = " [测试批注内容]"
new_run.append(t)

# 添加到段落
p_element.append(new_run)
print('添加了新的run')

# 保存
doc2.save('test_debug2_output.docx')
print('保存了文档')

# 重新加载检查
doc3 = Document('test_debug2_output.docx')
para3 = doc3.paragraphs[0]
print(f'\n输出文档段落: "{para3.text}"')
print(f'输出文档runs数: {len(para3.runs)}')
for i, run in enumerate(para3.runs):
    print(f'  Run {i}: "{run.text}" (长度: {len(run.text)})')
