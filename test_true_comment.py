"""
创建真正的Word批注 - 通过直接操作ZIP文件
"""
import zipfile
import os
import shutil
from xml.etree import ElementTree as ET
from xml.dom import minidom

# Word文档的命名空间
NAMESP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'rels': 'http://schemas.openxmlformats.org/package/2006/relationships'
}

def add_word_comment_to_run(docx_path, output_path, paragraph_index, run_index, comment_id, author, comment_text):
    """
    直接操作docx文件添加真正的Word批注

    Args:
        docx_path: 原始docx文件路径
        output_path: 输出文件路径
        paragraph_index: 段落索引
        run_index: run索引
        comment_id: 批注ID
        author: 作者
        comment_text: 批注内容
    """
    # 创建临时目录
    temp_dir = "temp_docx"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    # 解压docx文件
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    # 读取document.xml
    document_xml_path = os.path.join(temp_dir, "word", "document.xml")
    tree = ET.parse(document_xml_path)
    root = tree.getroot()

    # 注册命名空间
    for prefix, uri in NAMESP.items():
        ET.register_namespace(prefix, uri)

    # 找到目标段落和run
    # 在Word中，段落是w:p，run是w:r
    paragraphs = root.findall('.//w:ws:p', NAMESP)
    # 或者在没有w:ws的情况下
    if not paragraphs:
        paragraphs = root.findall('.//w:p', NAMESP)

    if paragraph_index >= len(paragraphs):
        print(f'段落索引超出范围: {paragraph_index} >= {len(paragraphs)}')
        shutil.rmtree(temp_dir)
        return False

    target_para = paragraphs[paragraph_index]
    runs = target_para.findall('.//w:r', NAMESP)

    if run_index >= len(runs):
        print(f'run索引超出范围: {run_index} >= {len(runs)}')
        shutil.rmtree(temp_dir)
        return False

    target_run = runs[run_index]

    # 为run添加背景色高亮
    rpr = target_run.find('w:rPr', NAMESP)
    if rpr is None:
        rpr = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        target_run.insert(0, rpr)

    shd = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFF00')
    rpr.append(shd)

    # 插入批注范围开始标记
    comment_start = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
    comment_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))

    # 找到run在段落中的位置
    para_children = list(target_para)
    run_position = para_children.index(target_run)

    # 插入批注标记
    target_para.insert(run_position, comment_start)

    # 插入批注引用
    comment_ref = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    rpr_ref = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    comment_ref.append(rpr_ref)
    ref_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}annotationRef')
    ref_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
    comment_ref.append(ref_elem)

    # 插入在run后面（由于已插入comment_start，位置+1）
    target_para.insert(run_position + 2, comment_ref)

    # 插入批注范围结束标记
    comment_end = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
    comment_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
    target_para.insert(run_position + 3, comment_end)

    # 保存修改后的document.xml
    tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)

    # 创建或更新comments.xml
    comments_xml_path = os.path.join(temp_dir, "word", "comments.xml")

    # 检查是否已存在comments.xml
    comments = None
    if os.path.exists(comments_xml_path):
        tree_comments = ET.parse(comments_xml_path)
        comments = tree_comments.getroot()
    else:
        # 创建新的comments.xml
        comments = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comments')

    # 创建批注
    comment = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
    comment.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
    comment.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', author)

    # 批注段落
    p = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    r = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    t = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = comment_text
    r.append(t)
    p.append(r)
    comment.append(p)

    comments.append(comment)

    # 保存comments.xml
    tree_comments = ET.ElementTree(comments)
    tree_comments.write(comments_xml_path, encoding='utf-8', xml_declaration=True)

    # 更新_rels文件以包含comments关系（如果还没有）
    rels_xml_path = os.path.join(temp_dir, "word", "_rels", "document.xml.rels")
    tree_rels = ET.parse(rels_xml_path)
    rels_root = tree_rels.getroot()

    # 检查是否已有comments关系
    has_comments_rel = False
    for child in rels_root:
        if 'comments' in str(child.get('Type', '')):
            has_comments_rel = True
            break

    if not has_comments_rel:
        # 添加comments关系
        new_rel = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
        new_rel.set('Id', 'rIdComments')
        new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
        new_rel.set('Target', 'comments.xml')
        rels_root.append(new_rel)

    tree_rels.write(rels_xml_path, encoding='utf-8', xml_declaration=True)

    # 重新打包zip文件
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(temp_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, temp_dir)
                zipf.write(file_path, arcname)

    # 清理临时目录
    shutil.rmtree(temp_dir)

    print(f'OK 批注已添加: 段落{paragraph_index}, run{run_index}, ID={comment_id}')
    return True

# 测试
if __name__ == '__main__':
    import sys
    sys.path.insert(0, r'f:\Agent项目\AI合同审查')

    # 创建测试文档
    from docx import Document
    doc = Document()
    doc.add_paragraph('这是第一个需要批注的文字。')
    doc.add_paragraph('这是第二个需要批注的文字。')
    doc.save('test_true_comment.docx')

    # 添加批注
    add_word_comment_to_run(
        'test_true_comment.docx',
        'test_true_comment_output.docx',
        0,  # 第一段
        0,  # 第一个run
        0,  # 批注ID
        'AI审核',
        '【高风险】这个条款存在问题，建议修改'
    )

    add_word_comment_to_run(
        'test_true_comment_output.docx',
        'test_true_comment_final.docx',
        1,  # 第二段
        0,  # 第一个run
        1,  # 批注ID
        'AI审核',
        '【中风险】这个条款需要优化'
    )

    print('\\nOK 完成！请打开 test_true_comment_final.docx 查看批注')
